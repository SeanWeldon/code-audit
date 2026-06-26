"""Generate a Bolder Apps branded Code Audit DOCX from the markdown audit doc.

Usage:
    python generate_audit_docx.py <AppName>-Code-Audit.md [output.docx]

Forked from code-takeover-estimation/assets/generate_review_docx.py with two
additive changes for the /full-code-audit deliverable:
  1. Renders an optional `## Code Review Scorecard` section (10-dimension
     table: Dimension | Grade | Risk | Evidence), with grade-colored cells.
  2. Doc-type-aware footer: if a scorecard is present the footer reads as a
     full audit; otherwise it falls back to the light-review footer (so this
     generator stays backward-compatible with takeover-audit-template docs).

Section headings drive output structure: ## Verdict, ## Code Review Scorecard,
## Findings Summary, ## Critical Findings (Blockers), ## What's Working Well,
## Recommended Path, ## Assumptions and Dependencies, ## Engagement Options.

Requires: python-docx
    uv venv && uv pip install python-docx
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# Bolder Apps brand tokens
ORANGE = "FF4626"
NAVY = "1A1A2E"
LIGHT_ORANGE = "FFF0ED"
LIGHT_GRAY = "F9F9F9"
WHITE = "FFFFFF"

# Severity colors (findings)
SEVERITY_COLORS = {
    "BLOCKER": "CC0000",   # red
    "HIGH": "CC6600",      # amber
    "MEDIUM": "A88300",    # yellow-gray
    "LOW": "666666",       # gray
}

# Grade colors (scorecard)
GRADE_COLORS = {
    "A": "2E7D32",   # green
    "B": "2E7D32",   # green
    "C": "CC6600",   # amber
    "D": "CC0000",   # red
    "F": "CC0000",   # red
}

# Verdict colors
VERDICT_COLORS = {
    "TAKEOVER VIABLE": "2E7D32",                  # green
    "VIABLE - STABILIZATION REQUIRED": "CC6600",  # amber
    "TAKEOVER NOT VIABLE": "CC0000",              # red
    "READY": "2E7D32",
    "NOT READY TO SHIP": "CC0000",
}


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.makeelement(qn("w:shd"), {qn("w:fill"): color_hex, qn("w:val"): "clear"})
    tc_pr.append(shading)


def write_cell(cell, text, bold=False, color=None, size=9, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    sizes = {1: 18, 2: 13, 3: 11}
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 11))
    run.font.color.rgb = RGBColor.from_string(NAVY)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def add_paragraph(doc, text, bold=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def add_verdict_block(doc, verdict_line, rationale):
    p = doc.add_paragraph()
    run = p.add_run(verdict_line)
    run.bold = True
    run.font.size = Pt(16)
    color = VERDICT_COLORS.get(verdict_line.strip(), NAVY)
    run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_after = Pt(6)

    if rationale:
        rp = doc.add_paragraph()
        rrun = rp.add_run(rationale)
        rrun.font.size = Pt(10)
        rp.paragraph_format.space_after = Pt(10)


def add_table(doc, headers, rows, color_col=None, color_map=None):
    """Generic branded table. If color_col is set, color that column by color_map
    (defaults to SEVERITY_COLORS). Keys are matched case-insensitively on the cell's
    leading token (so "A" matches grade "A", "BLOCKER" matches severity "BLOCKER")."""
    if color_map is None:
        color_map = SEVERITY_COLORS
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        write_cell(cell, h, bold=True, color=WHITE, size=9)
        set_cell_shading(cell, ORANGE)

    for r_idx, row in enumerate(rows):
        key = row[color_col].strip().upper() if color_col is not None and color_col < len(row) else None
        color = color_map.get(key) if key else None
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell_color = color if c_idx == color_col else None
            write_cell(cell, val, size=9, color=cell_color, bold=(c_idx == color_col and color is not None))
            if r_idx % 2 == 1 and c_idx != color_col:
                set_cell_shading(cell, LIGHT_GRAY)

    doc.add_paragraph("")
    return table


def parse_md_table(lines, start_idx):
    """Parse a pipe-delimited markdown table starting at start_idx. Returns (headers, rows, end_idx)."""
    if start_idx >= len(lines) or not lines[start_idx].strip().startswith("|"):
        return None, None, start_idx

    header_line = lines[start_idx].strip()
    headers = [c.strip() for c in header_line.strip("|").split("|")]

    sep_idx = start_idx + 1
    if sep_idx >= len(lines) or not re.match(r"^\s*\|[-:\s|]+\|\s*$", lines[sep_idx]):
        return None, None, start_idx

    rows = []
    i = sep_idx + 1
    while i < len(lines) and lines[i].strip().startswith("|"):
        row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if len(row) == len(headers):
            rows.append(row)
        i += 1
    return headers, rows, i


def parse_audit_md(md_text):
    """Parse the audit markdown into structured sections."""
    sections = {
        "title": "",
        "meta": [],          # list of "Key: Value" header lines
        "verdict": "",
        "rationale": "",
        "metrics": [],       # list of (key, value)
        "scorecard": [],     # list of [dimension, grade, risk, evidence]
        "findings": [],      # list of dicts {severity, area, finding, evidence}
        "critical_narrative": [],  # list of (title, body)
        "whats_working": [],       # list of bullet strings
        "recommended_path_body": "",
        "phased_plan": [],   # list of (phase, focus, hours, weeks)
        "assumptions": [],   # list of bullet strings
        "engagement_options": [],  # list of bullet strings
    }
    lines = md_text.splitlines()
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]

        if line.startswith("# ") and not sections["title"]:
            sections["title"] = line[2:].strip()
            i += 1
            continue

        if line.startswith("**") and ":" in line and not sections["verdict"]:
            sections["meta"].append(line.replace("**", "").strip())
            i += 1
            continue

        if line.strip().startswith("## Verdict"):
            i += 1
            while i < n and not lines[i].strip().startswith("##"):
                stripped = lines[i].strip()
                if stripped.startswith(">"):
                    sections["verdict"] = stripped.lstrip("> ").strip("*").strip()
                elif stripped.startswith("|"):
                    headers, rows, end = parse_md_table(lines, i)
                    if headers and rows:
                        sections["metrics"] = rows
                        i = end
                        continue
                elif stripped and not stripped.startswith("**") and not stripped.startswith("---"):
                    if not sections["rationale"]:
                        sections["rationale"] = stripped
                    else:
                        sections["rationale"] += " " + stripped
                i += 1
            continue

        if line.strip().startswith("## Code Review Scorecard"):
            i += 1
            while i < n and not lines[i].strip().startswith("## "):
                if lines[i].strip().startswith("|"):
                    headers, rows, end = parse_md_table(lines, i)
                    if headers and rows:
                        sections["scorecard"] = rows
                        i = end
                        continue
                i += 1
            continue

        if line.strip().startswith("## Findings Summary"):
            i += 1
            while i < n and not lines[i].strip().startswith("##"):
                if lines[i].strip().startswith("|"):
                    headers, rows, end = parse_md_table(lines, i)
                    if headers and rows:
                        for row in rows:
                            sections["findings"].append({
                                "severity": row[0],
                                "area": row[1] if len(row) > 1 else "",
                                "finding": row[2] if len(row) > 2 else "",
                                "evidence": row[3] if len(row) > 3 else "",
                            })
                        i = end
                        continue
                i += 1
            continue

        if line.strip().startswith("## Critical Findings"):
            i += 1
            current_title = None
            current_body = []
            while i < n and not lines[i].strip().startswith("## "):
                stripped = lines[i].strip()
                if stripped.startswith("### "):
                    if current_title:
                        sections["critical_narrative"].append((current_title, " ".join(current_body).strip()))
                    current_title = stripped[4:].strip()
                    current_body = []
                elif stripped and not stripped.startswith("---"):
                    current_body.append(stripped)
                i += 1
            if current_title:
                sections["critical_narrative"].append((current_title, " ".join(current_body).strip()))
            continue

        if line.strip().startswith("## What's Working Well"):
            i += 1
            while i < n and not lines[i].strip().startswith("## "):
                stripped = lines[i].strip()
                if stripped.startswith("- "):
                    sections["whats_working"].append(stripped[2:].strip())
                i += 1
            continue

        if line.strip().startswith("## Recommended Path"):
            i += 1
            body_buf = []
            while i < n and not lines[i].strip().startswith("## "):
                stripped = lines[i].strip()
                if stripped.startswith("|"):
                    headers, rows, end = parse_md_table(lines, i)
                    if headers and rows and "Phase" in headers[0]:
                        sections["phased_plan"] = rows
                        i = end
                        continue
                elif stripped and not stripped.startswith("###") and not stripped.startswith("---"):
                    body_buf.append(stripped)
                i += 1
            sections["recommended_path_body"] = " ".join(body_buf).strip()
            continue

        if line.strip().startswith("## Assumptions"):
            i += 1
            while i < n and not lines[i].strip().startswith("## "):
                stripped = lines[i].strip()
                if stripped.startswith("- "):
                    sections["assumptions"].append(stripped[2:].strip())
                i += 1
            continue

        if line.strip().startswith("## Engagement Options"):
            i += 1
            while i < n and not lines[i].strip().startswith("## "):
                stripped = lines[i].strip()
                if stripped.startswith("- "):
                    sections["engagement_options"].append(stripped[2:].strip())
                i += 1
            continue

        i += 1

    return sections


def render_docx(sections, output_path):
    doc = Document()

    p = doc.add_paragraph()
    run = p.add_run(sections["title"] or "Code Audit")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    for meta_line in sections["meta"]:
        mp = doc.add_paragraph()
        mp.add_run(meta_line).font.size = Pt(9)

    add_heading(doc, "Verdict", level=2)
    add_verdict_block(doc, sections["verdict"], sections["rationale"])

    if sections["metrics"]:
        add_heading(doc, "By the numbers", level=3)
        add_table(doc, ["Metric", "Value"], sections["metrics"])

    if sections["scorecard"]:
        add_heading(doc, "Code Review Scorecard", level=2)
        # Color the Grade column (index 1) by GRADE_COLORS.
        add_table(doc, ["Dimension", "Grade", "Risk", "Evidence"],
                  sections["scorecard"], color_col=1, color_map=GRADE_COLORS)

    if sections["findings"]:
        add_heading(doc, "Findings Summary", level=2)
        findings_rows = [[f["severity"], f["area"], f["finding"], f["evidence"]] for f in sections["findings"]]
        add_table(doc, ["Severity", "Area", "Finding", "Evidence"], findings_rows, color_col=0)

    if sections["critical_narrative"]:
        add_heading(doc, "Critical Findings (Blockers)", level=2)
        for title, body in sections["critical_narrative"]:
            add_heading(doc, title, level=3)
            add_paragraph(doc, body)

    if sections["whats_working"]:
        add_heading(doc, "What's Working Well", level=2)
        for bullet in sections["whats_working"]:
            p = doc.add_paragraph(bullet, style="List Bullet")
            p.runs[0].font.size = Pt(10)

    add_heading(doc, "Recommended Path", level=2)
    if sections["phased_plan"]:
        add_table(doc, ["Phase", "Focus", "Hours", "Weeks"], sections["phased_plan"])
    if sections["recommended_path_body"]:
        add_paragraph(doc, sections["recommended_path_body"])

    if sections["assumptions"]:
        add_heading(doc, "Assumptions and Dependencies", level=2)
        for bullet in sections["assumptions"]:
            p = doc.add_paragraph(bullet, style="List Bullet")
            p.runs[0].font.size = Pt(10)

    if sections["engagement_options"]:
        add_heading(doc, "Engagement Options", level=2)
        for bullet in sections["engagement_options"]:
            p = doc.add_paragraph(bullet, style="List Bullet")
            p.runs[0].font.size = Pt(10)

    # Doc-type-aware footer: scorecard present => full audit; else light-review fallback.
    footer = doc.add_paragraph()
    if sections["scorecard"]:
        footer_text = (
            "This is a comprehensive code audit produced with multi-agent review and "
            "independent cross-model verification. Remediation hours are priced separately "
            "in the accompanying estimation."
        )
    else:
        footer_text = (
            "This audit is a light review intended to determine takeover viability. "
            "Deeper architectural review, security audit, or compliance assessment "
            "may be appropriate as a follow-on engagement."
        )
    fr = footer.add_run(footer_text)
    fr.italic = True
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor.from_string("666666")
    footer.paragraph_format.space_before = Pt(18)

    doc.save(output_path)
    print(f"Wrote {output_path}")


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    md_path = Path(sys.argv[1])
    if not md_path.exists():
        print(f"Error: {md_path} does not exist", file=sys.stderr)
        sys.exit(1)

    if len(sys.argv) >= 3:
        output_path = Path(sys.argv[2])
    else:
        output_path = md_path.with_suffix(".docx")

    md_text = md_path.read_text()
    sections = parse_audit_md(md_text)
    render_docx(sections, output_path)


if __name__ == "__main__":
    main()
