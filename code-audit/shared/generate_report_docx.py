"""Generic Bolder-branded Markdown -> DOCX renderer for audit reports.

Unlike generate_audit_docx.py (which is hardwired to the /full-code-audit
scorecard structure and renders A-F letter grades), this renderer is
section-agnostic: it walks ANY markdown report and renders headings, paragraphs,
bullet lists, blockquotes, pipe tables, and embedded images (e.g. emulator
screenshots) into a branded DOCX. No letter grading, no scorecard assumptions.

It is the deep-code-audit companion: deep reports are severity-ranked (P0-P3),
not graded, and carry runtime screenshot evidence the scorecard generator cannot
embed.

Usage:
    python generate_report_docx.py <report.md> [output.docx]

Image paths in the markdown (![alt](path)) are resolved relative to the
markdown file's directory. Portrait screenshots are width-capped so they fit.

Requires: python-docx
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# Bolder Apps brand tokens
ORANGE = "FF4626"
NAVY = "1A1A2E"
LIGHT_GRAY = "F2F2F2"
WHITE = "FFFFFF"

# Severity colors (P0-P3). No letter grades by design.
SEVERITY_COLORS = {
    "P0": "CC0000",  # red
    "P1": "CC6600",  # amber
    "P2": "A88300",  # yellow-gray
    "P3": "666666",  # gray
}
SEV_RE = re.compile(r"^\**\s*(P[0-3])\b")
IMG_RE = re.compile(r"^!\[(.*?)\]\((.*?)\)\s*$")
HEADING_SIZES = {1: 18, 2: 14, 3: 11, 4: 10}


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {qn("w:fill"): color_hex, qn("w:val"): "clear"})
    tc_pr.append(shd)


def strip_inline(text):
    """Drop markdown emphasis/code/link syntax to plain text for run-splitting."""
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text)
    return text


def add_runs(paragraph, text, base_size=10, base_color=None):
    """Render inline **bold** / `code` into runs; everything else plain."""
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text)  # links -> label
    # split on bold and code, keeping delimiters
    parts = re.split(r"(\*\*.*?\*\*|`[^`]*`)", text)
    for part in parts:
        if not part:
            continue
        bold = code = False
        val = part
        if part.startswith("**") and part.endswith("**"):
            bold, val = True, part[2:-2]
        elif part.startswith("`") and part.endswith("`"):
            code, val = True, part[1:-1]
        run = paragraph.add_run(val)
        run.bold = bold
        run.font.size = Pt(base_size)
        if code:
            run.font.name = "Consolas"
        if base_color:
            run.font.color.rgb = RGBColor.from_string(base_color)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    run = p.add_run(strip_inline(text))
    run.bold = True
    run.font.size = Pt(HEADING_SIZES.get(level, 10))
    run.font.color.rgb = RGBColor.from_string(ORANGE if level == 1 else NAVY)
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(4)


def parse_table(lines, i):
    header = [c.strip() for c in lines[i].strip().strip("|").split("|")]
    sep = i + 1
    if sep >= len(lines) or not re.match(r"^\s*\|[-:\s|]+\|\s*$", lines[sep]):
        return None, None, i
    rows, j = [], sep + 1
    while j < len(lines) and lines[j].strip().startswith("|"):
        row = [c.strip() for c in lines[j].strip().strip("|").split("|")]
        if len(row) == len(header):
            rows.append(row)
        j += 1
    return header, rows, j


def sev_col(header):
    for idx, h in enumerate(header):
        if h.strip().lower() in ("sev", "severity"):
            return idx
    return None


def add_table(doc, header, rows):
    scol = sev_col(header)
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for c, h in enumerate(header):
        cell = table.rows[0].cells[c]
        cell.text = ""
        r = cell.paragraphs[0].add_run(strip_inline(h))
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(WHITE)
        set_cell_shading(cell, ORANGE)
    for ri, row in enumerate(rows):
        sev = None
        if scol is not None and scol < len(row):
            m = SEV_RE.match(row[scol])
            if m:
                sev = m.group(1)
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            color = SEVERITY_COLORS.get(sev) if ci == scol and sev else None
            add_runs(cell.paragraphs[0], val, base_size=8, base_color=color)
            if ci == scol and sev:
                cell.paragraphs[0].runs and setattr(cell.paragraphs[0].runs[0].font, "bold", True)
            elif ri % 2 == 1:
                set_cell_shading(cell, LIGHT_GRAY)
    doc.add_paragraph("")


def add_image(doc, path, base_dir):
    img = (base_dir / path).resolve()
    if not img.exists():
        add_runs(doc.add_paragraph(), f"[missing image: {path}]", base_color="CC0000")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.add_run().add_picture(str(img), width=Inches(2.3))
    except Exception as e:  # noqa: BLE001
        add_runs(doc.add_paragraph(), f"[image error {path}: {e}]", base_color="CC0000")


def render(md_path, out_path):
    base_dir = md_path.parent
    lines = md_path.read_text().splitlines()
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        s = line.strip()
        if not s:
            i += 1
            continue
        if s.startswith("|") and parse_table(lines, i)[0]:
            header, rows, j = parse_table(lines, i)
            add_table(doc, header, rows)
            i = j
            continue
        m = IMG_RE.match(s)
        if m:
            add_image(doc, m.group(2), base_dir)
            i += 1
            continue
        if s.startswith("#"):
            lvl = len(s) - len(s.lstrip("#"))
            add_heading(doc, s.lstrip("#").strip(), lvl)
            i += 1
            continue
        if s.startswith("---"):
            i += 1
            continue
        if s.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            r = p.add_run(strip_inline(s[2:]))
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor.from_string("555555")
            i += 1
            continue
        if re.match(r"^[-*] ", s):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, s[2:], base_size=10)
            i += 1
            continue
        if re.match(r"^\d+\. ", s):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\.\s*", "", s), base_size=10)
            i += 1
            continue
        add_runs(doc.add_paragraph(), s, base_size=10)
        i += 1

    footer = doc.add_paragraph()
    fr = footer.add_run(
        "Prepared by Bolder Apps. Findings are severity-ranked (P0 critical -> P3 low); "
        "P0/P1 items were independently verified. Remediation is scoped separately."
    )
    fr.italic = True
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor.from_string("666666")
    footer.paragraph_format.space_before = Pt(18)
    doc.save(out_path)
    print(f"Wrote {out_path}")


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    md_path = Path(sys.argv[1])
    if not md_path.exists():
        print(f"Error: {md_path} does not exist", file=sys.stderr)
        sys.exit(1)
    out_path = Path(sys.argv[2]) if len(sys.argv) >= 3 else md_path.with_suffix(".docx")
    render(md_path, out_path)


if __name__ == "__main__":
    main()
